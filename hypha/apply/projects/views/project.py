import copy
import datetime
import io
import json

from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import UserPassesTestMixin
from django.contrib.auth.models import Group
from django.core.exceptions import ObjectDoesNotExist, PermissionDenied
from django.db import transaction
from django.db.models import Q
from django.http import Http404, HttpResponse, HttpResponseRedirect
from django.shortcuts import get_object_or_404, render
from django.template.loader import get_template
from django.urls import reverse
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.utils.functional import cached_property
from django.utils.text import slugify
from django.utils.translation import gettext as _
from django.views import View
from django.views.generic import (
    CreateView,
    DetailView,
    UpdateView,
)
from django.views.generic.detail import SingleObjectMixin
from django_filters.views import FilterView
from django_htmx.http import (
    HttpResponseClientRefresh,
)
from django_tables2 import SingleTableMixin
from docx import Document
from htmldocx import HtmlToDocx
from rolepermissions.checkers import has_object_permission

from hypha.apply.activity.adapters.utils import get_users_for_groups
from hypha.apply.activity.messaging import MESSAGES, messenger
from hypha.apply.activity.models import ACTION, ALL, COMMENT, TEAM, Activity
from hypha.apply.activity.views import ActivityContextMixin, CommentFormView
from hypha.apply.stream_forms.models import BaseStreamForm
from hypha.apply.todo.options import (
    PAF_REQUIRED_CHANGES,
    PAF_WAITING_APPROVAL,
    PAF_WAITING_ASSIGNEE,
    PROJECT_SUBMIT_PAF,
    PROJECT_WAITING_CONTRACT,
    PROJECT_WAITING_CONTRACT_DOCUMENT,
    PROJECT_WAITING_CONTRACT_REVIEW,
    PROJECT_WAITING_INVOICE,
    PROJECT_WAITING_PF,
    PROJECT_WAITING_SOW,
)
from hypha.apply.todo.utils import get_project_lead_tasks
from hypha.apply.todo.views import (
    add_task_to_user,
    add_task_to_user_group,
    remove_tasks_for_user,
    remove_tasks_for_user_group,
)
from hypha.apply.users.decorators import (
    staff_or_finance_or_contracting_required,
    staff_required,
)
from hypha.apply.users.roles import CONTRACTING_GROUP_NAME
from hypha.apply.utils.models import PDFPageSettings
from hypha.apply.utils.pdfs import render_as_pdf
from hypha.apply.utils.storage import PrivateMediaView
from hypha.apply.utils.views import DelegateableView, DelegatedViewMixin, ViewDispatcher

from ...funds.files import generate_private_file_path
from ..filters import ProjectListFilter
from ..forms import (
    ApproveContractForm,
    ApproversForm,
    AssignApproversForm,
    ChangePAFStatusForm,
    ChangeProjectStatusForm,
    ProjectForm,
    ProjectSOWForm,
    SetPendingForm,
    SkipPAFApprovalProcessForm,
    SubmitContractDocumentsForm,
    UpdateProjectLeadForm,
    UpdateProjectTitleForm,
    UploadContractDocumentForm,
    UploadContractForm,
    UploadDocumentForm,
)
from ..models.project import (
    APPROVE,
    CONTRACTING,
    DRAFT,
    INTERNAL_APPROVAL,
    INVOICING_AND_REPORTING,
    PROJECT_PUBLIC_STATUSES,
    REQUEST_CHANGE,
    Contract,
    ContractDocumentCategory,
    ContractPacketFile,
    DocumentCategory,
    PacketFile,
    PAFApprovals,
    Project,
    ProjectSettings,
)
from ..permissions import has_permission
from ..tables import ProjectsListTable
from ..utils import (
    get_paf_status_display,
    get_placeholder_file,
    get_project_status_choices,
)
from .report import ReportingMixin


class ProjectGetObjectMixin:
    def get_object(self):
        return get_object_or_404(Project, submission__pk=self.kwargs["pk"])


@method_decorator(staff_required, name="dispatch")
class SendForApprovalView(View):
    form_class = SetPendingForm
    model = Project
    template_name = "application_projects/modals/send_for_approval.html"

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(Project, id=kwargs.get("pk"))
        # permission check
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.form_class(instance=self.object)
        project_settings = ProjectSettings.for_request(self.request)
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "project_settings": project_settings,
                "paf_approvals": PAFApprovals.objects.filter(project=self.object),
                "remaining_document_categories": DocumentCategory.objects.filter(
                    ~Q(packet_files__project=self.object)
                ),
                "object": self.object,
            },
        )

    def post(self, *args, **kwargs):
        form = self.form_class(
            self.request.POST, user=self.request.user, instance=self.object
        )
        project_settings = ProjectSettings.for_request(self.request)
        if form.is_valid():
            project = self.object
            old_stage = project.status
            form.save()

            # remove PAF submission task for staff group
            remove_tasks_for_user(
                code=PROJECT_SUBMIT_PAF,
                user=self.object.lead,
                related_obj=self.object,
            )

            # remove PAF rejection task for staff if exists
            remove_tasks_for_user(
                code=PAF_REQUIRED_CHANGES,
                user=self.object.lead,
                related_obj=self.object,
            )

            paf_approvals = self.object.paf_approvals.filter(approved=False)

            if project_settings.paf_approval_sequential:
                if paf_approvals:
                    user = paf_approvals.first().user
                    if user:
                        # notify only if first user/approver is updated
                        messenger(
                            MESSAGES.SEND_FOR_APPROVAL,
                            request=self.request,
                            user=self.request.user,
                            source=self.object,
                        )
                        messenger(
                            MESSAGES.APPROVE_PAF,
                            request=self.request,
                            user=self.request.user,
                            source=self.object,
                            related=[paf_approvals.first()],
                        )
                        # add PAF waiting approval task for paf_approval user
                        add_task_to_user(
                            code=PAF_WAITING_APPROVAL,
                            user=user,
                            related_obj=self.object,
                        )
                    else:
                        messenger(
                            MESSAGES.ASSIGN_PAF_APPROVER,
                            request=self.request,
                            user=self.request.user,
                            source=self.object,
                        )
                        # add PAF waiting assignee task for paf_approval reviewer_roles
                        add_task_to_user_group(
                            code=PAF_WAITING_ASSIGNEE,
                            user_group=paf_approvals.first().paf_reviewer_role.user_roles.all(),
                            related_obj=self.object,
                        )
            else:
                if paf_approvals.filter(user__isnull=False).exists():
                    messenger(
                        MESSAGES.SEND_FOR_APPROVAL,
                        request=self.request,
                        user=self.request.user,
                        source=self.object,
                    )
                    messenger(
                        MESSAGES.APPROVE_PAF,
                        request=self.request,
                        user=self.request.user,
                        source=self.object,
                        related=paf_approvals.filter(user__isnull=False),
                    )
                    # add PAF waiting approval task for paf_approvals users
                    for paf_approval in paf_approvals.filter(user__isnull=False):
                        add_task_to_user(
                            code=PAF_WAITING_APPROVAL,
                            user=paf_approval.user,
                            related_obj=self.object,
                        )
                if paf_approvals.filter(user__isnull=True).exists():
                    messenger(
                        MESSAGES.ASSIGN_PAF_APPROVER,
                        request=self.request,
                        user=self.request.user,
                        source=self.object,
                    )
                    # add PAF waiting assignee task for paf_approvals reviewer_roles
                    for paf_approval in paf_approvals.filter(user__isnull=True):
                        add_task_to_user_group(
                            code=PAF_WAITING_ASSIGNEE,
                            user_group=paf_approval.paf_reviewer_role.user_roles.all(),
                            related_obj=self.object,
                        )

            project.status = INTERNAL_APPROVAL
            project.save(update_fields=["status"])

            messenger(
                MESSAGES.PROJECT_TRANSITION,
                request=self.request,
                user=self.request.user,
                source=project,
                related=old_stage,
            )

            messages.success(self.request, _("PAF has been submitted for approval"))
            return HttpResponseClientRefresh()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "project_settings": project_settings,
                "paf_approvals": PAFApprovals.objects.filter(project=self.object),
                "remaining_document_categories": DocumentCategory.objects.filter(
                    ~Q(packet_files__project=self.object)
                ),
                "object": self.object,
            },
        )


# PROJECT DOCUMENTS
class UploadDocumentView(CreateView):
    form_class = UploadDocumentForm
    model = Project
    template_name = "application_projects/modals/supporting_documents_upload.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        self.category = get_object_or_404(
            DocumentCategory, id=kwargs.get("category_pk")
        )
        permission = has_object_permission(
            "upload_project_documents", request.user, obj=self.project
        )
        if not permission:
            raise PermissionDenied()
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        upload_document_form = self.form_class(
            instance=self.project, initial={"category": self.category}
        )
        return render(
            self.request,
            self.template_name,
            context={
                "form": upload_document_form,
                "value": _("Submit"),
                "category": self.category,
                "object": self.project,
            },
        )

    def post(self, request, *args, **kwargs):
        form = self.form_class(
            self.request.POST,
            request.FILES,
            instance=self.project,
            initial={"category": self.category},
        )
        if form.is_valid():
            form.instance.project = self.project
            form.save()
            messenger(
                MESSAGES.UPLOAD_DOCUMENT,
                request=self.request,
                user=self.request.user,
                source=self.project,
            )
            return HttpResponse(
                status=204,
                headers={
                    "HX-Trigger": json.dumps(
                        {
                            "supportingDocumentUpload": None,
                            "showMessage": _("Document has been uploaded"),
                        }
                    ),
                },
            )
        return render(
            self.request,
            self.template_name,
            context={"form": form, "value": _("Submit"), "object": self.project},
            status=400,
        )


@method_decorator(staff_required, name="dispatch")
class RemoveDocumentView(View):
    model = Project

    def delete(self, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        self.object = self.project.packet_files.get(pk=kwargs.get("document_pk"))
        self.object.delete()

        return HttpResponse(
            status=204,
            headers={
                "HX-Trigger": json.dumps(
                    {
                        "supportingDocumentRemove": None,
                        "showMessage": _("Document has been removed"),
                    }
                ),
            },
        )


@method_decorator(login_required, name="dispatch")
class RemoveContractDocumentView(View):
    model = Project

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        if not request.user.is_applicant or request.user != self.project.user:
            raise PermissionDenied
        return super().dispatch(request, *args, **kwargs)

    def delete(self, *args, **kwargs):
        self.object = self.project.contract_packet_files.get(
            pk=kwargs.get("document_pk")
        )
        self.object.delete()

        return HttpResponse(
            status=204,
            headers={
                "HX-Trigger": json.dumps(
                    {
                        "contractingDocumentRemove": None,
                        "showMessage": _("Contracting document has been removed"),
                    }
                ),
            },
        )


# GENERAL FORM VIEWS


@method_decorator(staff_required, name="dispatch")
class UpdateLeadView(View):
    model = Project
    form_class = UpdateProjectLeadForm
    template_name = "application_projects/modals/lead_update.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.form_class(instance=self.project)
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update"),
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        # Fetch the old lead from the database
        old_lead = copy.copy(self.project.lead)
        form = self.form_class(self.request.POST, instance=self.project)
        if form.is_valid():
            tasks = get_project_lead_tasks(self.project.lead, self.project)

            form.save()
            tasks.update(user=self.project.lead)

            messenger(
                MESSAGES.UPDATE_PROJECT_LEAD,
                request=self.request,
                user=self.request.user,
                source=self.project,
                related=old_lead or _("Unassigned"),
            )

            return HttpResponse(
                status=204,
                headers={
                    "HX-Trigger": json.dumps(
                        {
                            "leadUpdated": None,
                            "showMessage": _("Lead has been updated."),
                        }
                    ),
                },
            )
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update"),
                "object": self.project,
            },
        )


@method_decorator(staff_required, name="dispatch")
class UpdateProjectTitleView(DelegatedViewMixin, UpdateView):
    model = Project
    form_class = UpdateProjectTitleForm
    template_name = "application_projects/modals/project_title_update.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.form_class(instance=self.project)
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update"),
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        # Fetch the old lead from the database
        old_title = copy.copy(self.project.title)

        form = self.form_class(self.request.POST, instance=self.project)

        if form.is_valid():
            form.save()

            messenger(
                MESSAGES.UPDATE_PROJECT_TITLE,
                request=self.request,
                user=self.request.user,
                source=self.project,
                related=old_title,
            )

            return HttpResponse(
                status=204,
                headers={
                    "HX-Trigger": json.dumps(
                        {
                            "titleUpdated": None,
                            "showMessage": _("Title has been updated"),
                        }
                    ),
                },
            )
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update"),
                "object": self.project,
            },
        )


# CONTRACTS


class ContractsMixin:
    def get_context_data(self, **kwargs):
        project = self.get_object()
        contracts = project.contracts.select_related(
            "approver",
        ).order_by("-created_at")

        latest_contract = contracts.first()
        contract_to_approve = None
        contract_to_sign = None
        if latest_contract:
            if not latest_contract.signed_by_applicant:
                contract_to_sign = latest_contract
            elif not latest_contract.approver:
                contract_to_approve = latest_contract

        context = super().get_context_data(**kwargs)
        context["contract_to_approve"] = contract_to_approve
        context["contract_to_sign"] = contract_to_sign
        context["contracts"] = contracts.approved()
        context["contract"] = latest_contract
        return context


@method_decorator(staff_required, name="dispatch")
class ApproveContractView(View):
    form_class = ApproveContractForm
    model = Contract
    template_name = "application_projects/modals/approve_contract.html"

    def get_object(self):
        latest_contract = self.project.contracts.order_by("-created_at").first()
        if latest_contract and not latest_contract.approver:
            return latest_contract

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        self.object = self.get_object()
        # permission
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["instance"] = self.object
        kwargs.pop("user")
        return kwargs

    def get(self, *args, **kwargs):
        form = self.form_class(instance=self.object)
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Confirm"),
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        form = self.form_class(self.request.POST, instance=self.object)
        if form.is_valid():
            with transaction.atomic():
                form.instance.approver = self.request.user
                form.instance.approved_at = timezone.now()
                form.instance.project = self.project
                form.save()

                old_stage = self.project.status

                messenger(
                    MESSAGES.APPROVE_CONTRACT,
                    request=self.request,
                    user=self.request.user,
                    source=self.project,
                    related=self.object,
                )
                # remove Project waiting contract review task for staff
                remove_tasks_for_user(
                    code=PROJECT_WAITING_CONTRACT_REVIEW,
                    user=self.project.lead,
                    related_obj=self.project,
                )

                self.project.status = INVOICING_AND_REPORTING
                self.project.save(update_fields=["status"])

                messenger(
                    MESSAGES.PROJECT_TRANSITION,
                    request=self.request,
                    user=self.request.user,
                    source=self.project,
                    related=old_stage,
                )
                # add Project waiting invoice task for applicant
                add_task_to_user(
                    code=PROJECT_WAITING_INVOICE,
                    user=self.project.user,
                    related_obj=self.project,
                )

            messages.success(
                self.request,
                _(
                    "Contractor documents have been approved.  You can receive invoices from vendor now."
                ),
            )
            return HttpResponseClientRefresh()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Confirm"),
                "object": self.project,
            },
        )


@method_decorator(login_required, name="dispatch")
class UploadContractView(View):
    model = Project
    form_class = UploadContractForm
    template_name = "application_projects/modals/upload_contract.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        permission, _ = has_permission(
            "contract_upload", request.user, object=self.project
        )
        if permission:
            return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.get_form()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "user": self.request.user,
                "value": _("Upload") if self.request.user.is_applicant else _("Submit"),
                "object": self.project,
            },
        )

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs.pop("instance")
        kwargs.pop("user")
        return kwargs

    def get_form(self, *args, **kwargs):
        form = self.form_class(*args, **kwargs)
        if self.request.user.is_applicant:
            form.fields.pop("signed_and_approved")
        return form

    def post(self, *args, **kwargs):
        form = self.get_form(self.request.POST)

        if form.is_valid():
            if self.project.contracts.exists():
                form.instance = self.project.contracts.order_by("created_at").first()

            form.instance.project = self.project

            if self.request.user == self.project.user:
                form.instance.signed_by_applicant = True
                form.instance.uploaded_by_applicant_at = timezone.now()
                messages.success(self.request, _("Countersigned contract uploaded"))
            elif self.request.user.is_contracting or self.request.user.is_apply_staff:
                form.instance.uploaded_by_contractor_at = timezone.now()
                messages.success(self.request, _("Signed contract uploaded"))

            form.save()

            contract_signed_and_approved = form.cleaned_data.get("signed_and_approved")
            if contract_signed_and_approved:
                form.instance.approver = self.request.user
                form.instance.approved_at = timezone.now()
                form.instance.signed_and_approved = contract_signed_and_approved
                form.instance.signed_by_applicant = True
                form.instance.save(
                    update_fields=[
                        "approver",
                        "approved_at",
                        "signed_and_approved",
                        "signed_by_applicant",
                    ]
                )

                self.project.status = INVOICING_AND_REPORTING
                self.project.save(update_fields=["status"])
                old_stage = CONTRACTING

                messenger(
                    MESSAGES.PROJECT_TRANSITION,
                    request=self.request,
                    user=self.request.user,
                    source=self.project,
                    related=old_stage,
                )
                # remove Project waiting contract task for contracting/staff group
                if settings.STAFF_UPLOAD_CONTRACT:
                    remove_tasks_for_user(
                        code=PROJECT_WAITING_CONTRACT,
                        user=self.project.lead,
                        related_obj=self.project,
                    )
                else:
                    remove_tasks_for_user_group(
                        code=PROJECT_WAITING_CONTRACT,
                        user_group=Group.objects.filter(name=CONTRACTING_GROUP_NAME),
                        related_obj=self.project,
                    )
                # add Project waiting invoice task for applicant
                add_task_to_user(
                    code=PROJECT_WAITING_INVOICE,
                    user=self.project.user,
                    related_obj=self.project,
                )
            else:
                if self.request.user != self.project.user:
                    messenger(
                        MESSAGES.UPLOAD_CONTRACT,
                        request=self.request,
                        user=self.request.user,
                        source=self.project,
                        related=form.instance,
                    )
                    # remove Project waiting contract task for contracting/staff group
                    if settings.STAFF_UPLOAD_CONTRACT:
                        remove_tasks_for_user(
                            code=PROJECT_WAITING_CONTRACT,
                            user=self.project.lead,
                            related_obj=self.project,
                        )
                    else:
                        remove_tasks_for_user_group(
                            code=PROJECT_WAITING_CONTRACT,
                            user_group=Group.objects.filter(
                                name=CONTRACTING_GROUP_NAME
                            ),
                            related_obj=self.project,
                        )
                    # add Project waiting contract document task for applicant
                    add_task_to_user(
                        code=PROJECT_WAITING_CONTRACT_DOCUMENT,
                        user=self.project.user,
                        related_obj=self.project,
                    )
            return HttpResponseClientRefresh()

        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "user": self.request.user,
                "value": _("Upload") if self.request.user.is_applicant else _("Submit"),
                "object": self.project,
            },
        )


class SkipPAFApprovalProcessView(UpdateView):
    model = Project
    form_class = SkipPAFApprovalProcessForm

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(Project, id=kwargs.get("pk"))
        # permissions
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        project = self.object
        old_stage = project.status
        project.is_locked = True
        project.status = CONTRACTING
        super().form_valid(form)

        # remove PAF submission task for staff group
        remove_tasks_for_user(
            code=PROJECT_SUBMIT_PAF,
            user=self.object.lead,
            related_obj=self.object,
        )

        # remove PAF rejection task for staff if exists
        remove_tasks_for_user(
            code=PAF_REQUIRED_CHANGES,
            user=self.object.lead,
            related_obj=self.object,
        )

        messenger(
            MESSAGES.PROJECT_TRANSITION,
            request=self.request,
            user=self.request.user,
            source=self.object,
            related=old_stage,
        )
        # add project waiting contract task to staff/contracting groups
        if settings.STAFF_UPLOAD_CONTRACT:
            add_task_to_user(
                code=PROJECT_WAITING_CONTRACT,
                user=self.object.lead,
                related_obj=self.object,
            )
        else:
            add_task_to_user_group(
                code=PROJECT_WAITING_CONTRACT,
                user_group=Group.objects.filter(name=CONTRACTING_GROUP_NAME),
                related_obj=self.object,
            )
        return HttpResponseClientRefresh()


@method_decorator(login_required, name="dispatch")
class SubmitContractDocumentsView(View):
    model = Project
    form_class = SubmitContractDocumentsForm
    template_name = "application_projects/modals/submit_contracting_documents.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        if ContractDocumentCategory.objects.filter(
            ~Q(contract_packet_files__project=self.project) & Q(required=True)
        ).exists():
            raise PermissionDenied
        contract = self.project.contracts.order_by("-created_at").first()
        permission, _ = has_permission(
            "submit_contract_documents",
            request.user,
            object=self.project,
            raise_exception=True,
            contract=contract,
        )
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.form_class()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Submit"),
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        form = self.form_class(self.request.POST, instance=self.project)
        if form.is_valid():
            form.save()

            self.project.submitted_contract_documents = True
            self.project.save(update_fields=["submitted_contract_documents"])

            messenger(
                MESSAGES.SUBMIT_CONTRACT_DOCUMENTS,
                request=self.request,
                user=self.request.user,
                source=self.project,
            )
            # remove project waiting contract documents task for applicant
            remove_tasks_for_user(
                code=PROJECT_WAITING_CONTRACT_DOCUMENT,
                user=self.project.user,
                related_obj=self.project,
            )
            # add project waiting contract review task for staff
            add_task_to_user(
                code=PROJECT_WAITING_CONTRACT_REVIEW,
                user=self.project.lead,
                related_obj=self.project,
            )

            messages.success(self.request, _("Contract documents submitted"))
            return HttpResponseClientRefresh()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Submit"),
                "object": self.project,
            },
        )


@method_decorator(login_required, name="dispatch")
class UploadContractDocumentView(View):
    form_class = UploadContractDocumentForm
    model = Project
    context_name = "contract_document_form"
    template_name = "application_projects/modals/contracting_documents_upload.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, id=kwargs.get("pk"))
        self.category = get_object_or_404(
            ContractDocumentCategory, id=kwargs.get("category_pk")
        )
        if request.user != self.project.user or not request.user.is_applicant:
            raise PermissionDenied
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        upload_contract_document_form = self.form_class(
            instance=self.project, initial={"category": self.category}
        )
        return render(
            self.request,
            self.template_name,
            context={
                "form": upload_contract_document_form,
                "value": _("Submit"),
                "category": self.category,
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        form = self.form_class(
            self.request.POST,
            self.request.FILES,
            instance=self.project,
            initial={"category": self.category},
        )
        if form.is_valid():
            form.instance.project = self.project
            form.save()

            return HttpResponse(
                status=204,
                headers={
                    "HX-Trigger": json.dumps(
                        {
                            "contractingDocumentUpload": None,
                            "showMessage": _("Contracting document has been uploaded"),
                        }
                    ),
                },
            )
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Submit"),
                "category": self.category,
                "object": self.project,
            },
        )


# PROJECT VIEW


@method_decorator(login_required, name="dispatch")
class ChangePAFStatusView(View):
    form_class = ChangePAFStatusForm
    model = Project
    template_name = "application_projects/modals/pafstatus_update.html"

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(Project, pk=self.kwargs["pk"])
        permission, _ = has_permission(
            "paf_status_update",
            self.request.user,
            object=self.object,
            raise_exception=True,
            request=request,
        )
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.form_class(instance=self.object)
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update Status"),
                "object": self.object,
            },
        )

    def post(self, *args, **kwargs):
        form = self.form_class(self.request.POST, instance=self.object)
        if form.is_valid():
            form.save()
            project_settings = ProjectSettings.for_request(self.request)
            paf_approval = self.request.user.paf_approvals.filter(
                project=self.object, approved=False
            ).first()
            if not paf_approval:
                # get paf project form for not-assigned case
                if project_settings.paf_approval_sequential:
                    paf_approval = self.object.paf_approvals.filter(
                        approved=False
                    ).first()
                else:
                    for approval in self.object.paf_approvals.filter(approved=False):
                        if self.request.user.id in [
                            role_user.id
                            for role_user in get_users_for_groups(
                                list(approval.paf_reviewer_role.user_roles.all()),
                                exact_match=True,
                            )
                        ]:
                            paf_approval = approval
                            break
                    else:
                        # should never be the case but still to avoid 500.
                        raise PermissionDenied(
                            "User don't have project form approver roles"
                        )

            paf_status = form.cleaned_data.get("paf_status")
            comment = form.cleaned_data.get("comment", "")

            paf_status_update_message = _(
                "<p>{role} has updated project form status to {paf_status}.</p>"
            ).format(
                role=paf_approval.paf_reviewer_role.label,
                paf_status=get_paf_status_display(paf_status).lower(),
            )
            Activity.objects.create(
                user=self.request.user,
                type=ACTION,
                source=self.object,
                timestamp=timezone.now(),
                message=paf_status_update_message,
                visibility=TEAM,
            )

            if paf_status == REQUEST_CHANGE:
                old_stage = self.object.status
                self.object.status = DRAFT
                self.object.save(update_fields=["status"])
                paf_approval.save()

                # remove PAF waiting assignee/approval task for paf approval user/reviewer roles.
                if project_settings.paf_approval_sequential:
                    if paf_approval.user:
                        remove_tasks_for_user(
                            code=PAF_WAITING_APPROVAL,
                            user=paf_approval.user,
                            related_obj=self.object,
                        )
                    else:
                        remove_tasks_for_user_group(
                            code=PAF_WAITING_ASSIGNEE,
                            user_group=paf_approval.paf_reviewer_role.user_roles.all(),
                            related_obj=self.object,
                        )
                else:
                    for approval in self.object.paf_approvals.filter(approved=False):
                        if approval.user:
                            remove_tasks_for_user(
                                code=PAF_WAITING_APPROVAL,
                                user=approval.user,
                                related_obj=self.object,
                            )
                        else:
                            remove_tasks_for_user_group(
                                code=PAF_WAITING_ASSIGNEE,
                                user_group=approval.paf_reviewer_role.user_roles.all(),
                                related_obj=self.object,
                            )

                if not paf_approval.user:
                    paf_approval.user = self.request.user
                    paf_approval.save(update_fields=["user"])

                messenger(
                    MESSAGES.REQUEST_PROJECT_CHANGE,
                    request=self.request,
                    user=self.request.user,
                    source=self.object,
                    comment=comment,
                )
                messenger(
                    MESSAGES.PROJECT_TRANSITION,
                    request=self.request,
                    user=self.request.user,
                    source=self.object,
                    related=old_stage,
                )
                # add PAF required changes task to staff user group
                add_task_to_user(
                    code=PAF_REQUIRED_CHANGES,
                    user=self.object.lead,
                    related_obj=self.object,
                )

                messages.success(
                    self.request, _("Project form status has been updated")
                )
            elif paf_status == APPROVE:
                # remove task for paf approval user/user_group related to this paf_approval of project
                if paf_approval.user:
                    remove_tasks_for_user(
                        code=PAF_WAITING_APPROVAL,
                        user=paf_approval.user,
                        related_obj=self.object,
                    )
                else:
                    remove_tasks_for_user_group(
                        code=PAF_WAITING_ASSIGNEE,
                        user_group=paf_approval.paf_reviewer_role.user_roles.all(),
                        related_obj=self.object,
                    )
                paf_approval.approved = True
                paf_approval.approved_at = timezone.now()
                paf_approval.user = self.request.user
                paf_approval.save(update_fields=["approved", "approved_at", "user"])
                if project_settings.paf_approval_sequential:
                    # notify next approver
                    if self.object.paf_approvals.filter(approved=False).exists():
                        next_paf_approval = self.object.paf_approvals.filter(
                            approved=False
                        ).first()
                        if next_paf_approval.user:
                            messenger(
                                MESSAGES.APPROVE_PAF,
                                request=self.request,
                                user=self.request.user,
                                source=self.object,
                                related=[next_paf_approval],
                            )
                            # add PAF waiting approval task for next paf approval user
                            add_task_to_user(
                                code=PAF_WAITING_APPROVAL,
                                user=next_paf_approval.user,
                                related_obj=self.object,
                            )
                        else:
                            messenger(
                                MESSAGES.ASSIGN_PAF_APPROVER,
                                request=self.request,
                                user=self.request.user,
                                source=self.object,
                            )
                            # add PAF waiting assignee task for nex paf approval reviewer roles
                            add_task_to_user_group(
                                code=PAF_WAITING_ASSIGNEE,
                                user_group=next_paf_approval.paf_reviewer_role.user_roles.all(),
                                related_obj=self.object,
                            )
                messages.success(self.request, _("Project form has been approved"))

            if form.cleaned_data["comment"]:
                comment = f'<p>"{form.cleaned_data["comment"]}."</p>'

                message = paf_status_update_message + comment

                Activity.objects.create(
                    user=self.request.user,
                    type=COMMENT,
                    source=self.object,
                    timestamp=timezone.now(),
                    message=message,
                    visibility=TEAM,
                )

            if self.object.is_approved_by_all_paf_reviewers:
                old_stage = self.object.status
                self.object.is_locked = True
                self.object.status = CONTRACTING
                self.object.save(update_fields=["is_locked", "status"])

                messenger(
                    MESSAGES.PROJECT_TRANSITION,
                    request=self.request,
                    user=self.request.user,
                    source=self.object,
                    related=old_stage,
                )
                # add project waiting contract task to staff/contracting groups
                if settings.STAFF_UPLOAD_CONTRACT:
                    add_task_to_user(
                        code=PROJECT_WAITING_CONTRACT,
                        user=self.object.lead,
                        related_obj=self.object,
                    )
                else:
                    add_task_to_user_group(
                        code=PROJECT_WAITING_CONTRACT,
                        user_group=Group.objects.filter(name=CONTRACTING_GROUP_NAME),
                        related_obj=self.object,
                    )
            return HttpResponseClientRefresh()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update Status"),
                "object": self.object,
            },
        )


class ChangeProjectstatusView(View):
    """
    Project status can be updated manually only in 'IN PROGRESS, CLOSING and COMPLETE' state.
    """

    form_class = ChangeProjectStatusForm
    model = Project
    template_name = "application_projects/modals/project_status_update.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        permission, _ = has_permission(
            "project_status_update", request.user, self.project, raise_exception=True
        )
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.form_class(instance=self.project)
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update"),
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        old_stage = self.project.status
        form = self.form_class(self.request.POST, instance=self.project)

        if form.is_valid():
            form.save()

            comment = form.cleaned_data.get("comment", "")

            if comment:
                Activity.objects.create(
                    user=self.request.user,
                    type=COMMENT,
                    source=self.project,
                    timestamp=timezone.now(),
                    message=comment,
                    visibility=ALL,
                )

            messenger(
                MESSAGES.PROJECT_TRANSITION,
                request=self.request,
                user=self.request.user,
                source=self.project,
                related=old_stage,
            )

            messages.success(self.request, _("Project status has been updated"))
            return HttpResponseClientRefresh()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Update"),
                "object": self.project,
            },
        )


@method_decorator(login_required, name="dispatch")
class UpdateAssignApproversView(View):
    form_class = AssignApproversForm
    model = Project
    template_name = "application_projects/modals/assign_pafapprovers.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        permission, _ = has_permission(
            "update_paf_assigned_approvers",
            request.user,
            self.project,
            raise_exception=True,
            request=request,
        )
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        from ..forms.project import get_latest_project_paf_approval_via_roles

        form = self.form_class(user=self.request.user, instance=self.project)
        paf_approval = get_latest_project_paf_approval_via_roles(
            project=self.project, roles=self.request.user.groups.all()
        )
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Submit"),
                "pafapprover_exists": True if paf_approval.user else False,
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        from ..forms.project import get_latest_project_paf_approval_via_roles

        form = self.form_class(
            self.request.POST, user=self.request.user, instance=self.project
        )

        old_paf_approval = get_latest_project_paf_approval_via_roles(
            project=self.project, roles=self.request.user.groups.all()
        )

        if form.is_valid():
            form.save()

            # remove current task of user/user_group related to latest paf_approval of project
            if old_paf_approval.user:
                remove_tasks_for_user(
                    code=PAF_WAITING_APPROVAL,
                    user=old_paf_approval.user,
                    related_obj=self.project,
                )
            else:
                remove_tasks_for_user_group(
                    code=PAF_WAITING_ASSIGNEE,
                    user_group=old_paf_approval.paf_reviewer_role.user_roles.all(),
                    related_obj=self.project,
                )

            paf_approval = get_latest_project_paf_approval_via_roles(
                project=self.project, roles=self.request.user.groups.all()
            )

            if paf_approval.user:
                messenger(
                    MESSAGES.APPROVE_PAF,
                    request=self.request,
                    user=self.request.user,
                    source=self.project,
                    related=[paf_approval],
                )
                # add PAF waiting approval task to updated paf_approval user
                add_task_to_user(
                    code=PAF_WAITING_APPROVAL,
                    user=paf_approval.user,
                    related_obj=self.project,
                )
            else:
                messenger(
                    MESSAGES.ASSIGN_PAF_APPROVER,
                    request=self.request,
                    user=self.request.user,
                    source=self.project,
                )
                # add paf waiting for assignee task
                add_task_to_user_group(
                    code=PAF_WAITING_ASSIGNEE,
                    user_group=paf_approval.paf_reviewer_role.user_roles.all(),
                    related_obj=self.project,
                )

            return HttpResponseClientRefresh()

        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Submit"),
                "pafapprover_exists": True if old_paf_approval.user else False,
                "object": self.project,
            },
        )


class UpdatePAFApproversView(View):
    form_class = ApproversForm
    model = Project
    template_name = "application_projects/modals/update_pafapprovers.html"

    def dispatch(self, request, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        permission, _ = has_permission(
            "paf_approvers_update",
            request.user,
            self.project,
            raise_exception=True,
            request=request,
        )
        return super().dispatch(request, *args, **kwargs)

    def get(self, *args, **kwargs):
        form = self.form_class(instance=self.project)
        project_settings = ProjectSettings.for_request(self.request)
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Submit"),
                "project_settings": project_settings,
                "object": self.project,
            },
        )

    def post(self, *args, **kwargs):
        form = self.form_class(self.request.POST, instance=self.project)

        project_settings = ProjectSettings.for_request(self.request)
        if form.is_valid():
            old_approvers = None
            if self.project.paf_approvals.exists():
                old_approvers = list(
                    self.project.paf_approvals.filter(approved=False).values_list(
                        "user__id", flat=True
                    )
                )
            # remove PAF waiting assignee/approval task for paf approval user/reviewer roles.
            if project_settings.paf_approval_sequential:
                paf_approval = self.project.paf_approvals.filter(approved=False).first()
                if paf_approval.user:
                    remove_tasks_for_user(
                        code=PAF_WAITING_APPROVAL,
                        user=paf_approval.user,
                        related_obj=self.project,
                    )
                else:
                    remove_tasks_for_user_group(
                        code=PAF_WAITING_ASSIGNEE,
                        user_group=paf_approval.paf_reviewer_role.user_roles.all(),
                        related_obj=self.project,
                    )
            else:
                for approval in self.project.paf_approvals.filter(approved=False):
                    if approval.user:
                        remove_tasks_for_user(
                            code=PAF_WAITING_APPROVAL,
                            user=approval.user,
                            related_obj=self.project,
                        )
                    else:
                        remove_tasks_for_user_group(
                            code=PAF_WAITING_ASSIGNEE,
                            user_group=approval.paf_reviewer_role.user_roles.all(),
                            related_obj=self.project,
                        )

            form.save()

            paf_approvals = self.project.paf_approvals.filter(approved=False)

            if old_approvers and paf_approvals:
                # if approvers exists already
                if project_settings.paf_approval_sequential:
                    user = paf_approvals.first().user
                    if user and user.id != old_approvers[0]:
                        # notify only if first user/approver is updated
                        messenger(
                            MESSAGES.APPROVE_PAF,
                            request=self.request,
                            user=self.request.user,
                            source=self.project,
                            related=[paf_approvals.first()],
                        )
                        # add PAF waiting approval task to paf_approval user
                        add_task_to_user(
                            code=PAF_WAITING_APPROVAL,
                            user=user,
                            related_obj=self.project,
                        )
                    elif not user:
                        messenger(
                            MESSAGES.ASSIGN_PAF_APPROVER,
                            request=self.request,
                            user=self.request.user,
                            source=self.project,
                        )
                        # add PAF waiting assignee to paf_approvals reviewer roles
                        add_task_to_user_group(
                            code=PAF_WAITING_ASSIGNEE,
                            user_group=paf_approvals.first().paf_reviewer_role.user_roles.all(),
                            related_obj=self.project,
                        )
                else:
                    if paf_approvals.filter(user__isnull=False).exists():
                        messenger(
                            MESSAGES.APPROVE_PAF,
                            request=self.request,
                            user=self.request.user,
                            source=self.project,
                            related=paf_approvals.filter(user__isnull=False),
                        )
                        # add PAF waiting approval task for paf_approvals users
                        for paf_approval in paf_approvals.filter(user__isnull=False):
                            add_task_to_user(
                                code=PAF_WAITING_APPROVAL,
                                user=paf_approval.user,
                                related_obj=self.project,
                            )
                    if paf_approvals.filter(user__isnull=True).exists():
                        messenger(
                            MESSAGES.ASSIGN_PAF_APPROVER,
                            request=self.request,
                            user=self.request.user,
                            source=self.project,
                        )
                        # add PAF waiting assignee task for paf_approvals reviewer_roles
                        for paf_approval in paf_approvals.filter(user__isnull=True):
                            add_task_to_user_group(
                                code=PAF_WAITING_ASSIGNEE,
                                user_group=paf_approval.paf_reviewer_role.user_roles.all(),
                                related_obj=self.project,
                            )
            elif paf_approvals:
                # :todo: check if this is covering any case(might be a duplicate of SendForApprovalView)
                if paf_approvals.filter(user__isnull=False).exists():
                    messenger(
                        MESSAGES.APPROVE_PAF,
                        request=self.request,
                        user=self.request.user,
                        source=self.project,
                        related=paf_approvals.filter(user__isnull=False),
                    )
                    # add PAF waiting approval task for paf_approvals users
                    for paf_approval in paf_approvals.filter(user__isnull=False):
                        add_task_to_user(
                            code=PAF_WAITING_APPROVAL,
                            user=paf_approval.user,
                            related_obj=self.project,
                        )
                if paf_approvals.filter(user__isnull=True).exists():
                    messenger(
                        MESSAGES.ASSIGN_PAF_APPROVER,
                        request=self.request,
                        user=self.request.user,
                        source=self.project,
                    )
                    # add PAF waiting assignee task for paf_approvals reviewer_roles
                    for paf_approval in paf_approvals.filter(user__isnull=True):
                        add_task_to_user_group(
                            code=PAF_WAITING_ASSIGNEE,
                            user_group=paf_approval.paf_reviewer_role.user_roles.all(),
                            related_obj=self.project,
                        )

            messages.success(
                self.request, _("Project form approvers have been updated")
            )
            return HttpResponseClientRefresh()
        return render(
            self.request,
            self.template_name,
            context={
                "form": form,
                "value": _("Submit"),
                "project_settings": project_settings,
                "object": self.project,
            },
        )


class BaseProjectDetailView(ReportingMixin, ProjectGetObjectMixin, DetailView):
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["statuses"] = get_project_status_choices()
        context["current_status_index"] = [
            status for status, _ in get_project_status_choices()
        ].index(self.object.status)
        context["supporting_documents_configured"] = (
            True if DocumentCategory.objects.count() else False
        )
        context["contracting_documents_configured"] = (
            True if ContractDocumentCategory.objects.count() else False
        )
        return context


class AdminProjectDetailView(
    ActivityContextMixin,
    DelegateableView,
    ContractsMixin,
    BaseProjectDetailView,
):
    form_views = [
        CommentFormView,
    ]
    model = Project
    template_name_suffix = "_admin_detail"

    def dispatch(self, *args, **kwargs):
        project = self.get_object()
        permission, _ = has_permission(
            "project_access", self.request.user, object=project, raise_exception=True
        )
        return super().dispatch(*args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["statuses"] = get_project_status_choices()
        context["current_status_index"] = [
            status for status, _ in get_project_status_choices()
        ].index(self.object.status)
        project_settings = ProjectSettings.for_request(self.request)
        context["project_settings"] = project_settings
        context["paf_approvals"] = PAFApprovals.objects.filter(project=self.object)
        return context


class ApplicantProjectDetailView(
    ActivityContextMixin,
    DelegateableView,
    ContractsMixin,
    BaseProjectDetailView,
):
    form_views = [
        CommentFormView,
    ]

    model = Project
    template_name_suffix = "_detail"

    def dispatch(self, request, *args, **kwargs):
        project = self.get_object()
        permission, _ = has_permission(
            "project_access", request.user, object=project, raise_exception=True
        )
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["statuses"] = PROJECT_PUBLIC_STATUSES
        context["current_status_index"] = [
            status for status, _ in PROJECT_PUBLIC_STATUSES
        ].index(self.object.status)
        context["remaining_contract_document_categories"] = (
            ContractDocumentCategory.objects.filter(
                ~Q(contract_packet_files__project=self.object)
            )
        )
        return context


class ProjectDetailView(ViewDispatcher):
    admin_view = AdminProjectDetailView
    finance_view = AdminProjectDetailView
    contracting_view = AdminProjectDetailView
    applicant_view = ApplicantProjectDetailView


@method_decorator(login_required, name="dispatch")
class ProjectPrivateMediaView(UserPassesTestMixin, PrivateMediaView):
    """
    See also hypha/apply/funds/files.py
    """

    raise_exception = True

    def dispatch(self, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        return super().dispatch(*args, **kwargs)

    def get_media(self, *args, **kwargs):
        if "file_pk" in kwargs:
            document = PacketFile.objects.get(pk=kwargs["file_pk"])
            if document.project != self.project:
                raise Http404
            return document.document
        else:
            field_id = kwargs["field_id"]
            file_name = kwargs["file_name"]
            path_to_file = generate_private_file_path(
                self.project.pk, field_id, file_name, path_start="project"
            )
            return self.storage.open(path_to_file)

    def test_func(self):
        if self.request.user.is_org_faculty:
            return True

        if self.request.user == self.project.user:
            return True

        return False


@method_decorator(login_required, name="dispatch")
class CategoryTemplatePrivateMediaView(PrivateMediaView):
    raise_exception = True

    def dispatch(self, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        self.category_type = kwargs["type"]
        permission, _ = has_permission(
            "project_access",
            self.request.user,
            object=self.project,
            raise_exception=True,
        )
        return super().dispatch(*args, **kwargs)

    def get_media(self, *args, **kwargs):
        if self.category_type == "project_document":
            category = DocumentCategory.objects.get(pk=kwargs["category_pk"])
        elif self.category_type == "contract_document":
            category = ContractDocumentCategory.objects.get(pk=kwargs["category_pk"])
        else:
            raise Http404
        if not category.template:
            raise Http404
        return category.template


@method_decorator(login_required, name="dispatch")
class ContractPrivateMediaView(UserPassesTestMixin, PrivateMediaView):
    raise_exception = True

    def dispatch(self, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        return super().dispatch(*args, **kwargs)

    def get_media(self, *args, **kwargs):
        document = Contract.objects.get(pk=kwargs["file_pk"])
        if document.project != self.project:
            raise Http404
        return document.file

    def test_func(self):
        if self.request.user.is_apply_staff or self.request.user.is_contracting:
            return True

        if self.request.user == self.project.user:
            return True

        return False


@method_decorator(login_required, name="dispatch")
class ContractDocumentPrivateMediaView(PrivateMediaView):
    raise_exception = True

    def dispatch(self, *args, **kwargs):
        self.project = get_object_or_404(Project, submission__id=self.kwargs["pk"])
        self.document = ContractPacketFile.objects.get(pk=kwargs["file_pk"])
        permission, _ = has_permission(
            "view_contract_documents",
            self.request.user,
            object=self.project,
            contract_category=self.document.category,
            raise_exception=True,
        )
        return super().dispatch(*args, **kwargs)

    def get_media(self, *args, **kwargs):
        if self.document.project != self.project:
            raise Http404
        return self.document.document


# PROJECT FORM VIEWS


@method_decorator(staff_or_finance_or_contracting_required, name="dispatch")
class ProjectDetailApprovalView(DetailView):
    model = Project
    template_name_suffix = "_approval_detail"


@method_decorator(staff_or_finance_or_contracting_required, name="dispatch")
class ProjectSOWView(DetailView):
    model = Project
    template_name_suffix = "_sow_detail"


@method_decorator(staff_or_finance_or_contracting_required, name="dispatch")
class ProjectSOWDownloadView(SingleObjectMixin, View):
    model = Project

    def get(self, request, *args, **kwargs):
        export_type = kwargs.get("export_type", "pdf")
        self.object = self.get_object()
        context = {}
        context["sow_data"] = self.get_sow_data_with_field(self.object)
        context["org_name"] = settings.ORG_LONG_NAME
        context["id"] = self.object.id
        context["title"] = self.object.title
        context["project_link"] = self.request.build_absolute_uri(
            object.get_absolute_url()
        )
        template_path = "application_projects/sow_export.html"

        if export_type == "pdf":
            pdf_page_settings = PDFPageSettings.load(request_or_site=request)
            context["pagesize"] = pdf_page_settings.download_page_size
            context["show_footer"] = True
            return render_as_pdf(
                request=request,
                template_name=template_path,
                context=context,
                filename=self.get_slugified_file_name(export_type),
            )
        elif export_type == "docx":
            context["show_footer"] = False

            return self.render_as_docx(
                context=context,
                template=get_template(template_path),
                filename=self.get_slugified_file_name(export_type),
            )
        else:
            raise Http404(f"{export_type} type not supported at the moment")

    def render_as_docx(self, context, template, filename):
        html = template.render(context)

        buf = io.BytesIO()
        document = Document()
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html, document)
        document.save(buf)

        response = HttpResponse(buf.getvalue(), content_type="application/docx")
        response["Content-Disposition"] = f"attachment; filename={filename}"
        return response

    def get_slugified_file_name(self, export_type):
        return f"{datetime.date.today().strftime('%Y%m%d')}-{slugify(self.object.title)}.{export_type}"

    def get_sow_data_with_field(self, project):
        data_dict = {}
        if project.submission.page.specific.sow_forms.exists() and hasattr(
            project, "sow"
        ):
            form_data_dict = project.sow.form_data
            for field in project.sow.form_fields.raw_data:
                if field.get("type", None) in ["file", "multi_file"]:
                    continue
                if field["id"] in form_data_dict.keys():
                    if (
                        isinstance(field["value"], dict)
                        and "field_label" in field["value"]
                    ):
                        data_dict[field["value"]["field_label"]] = form_data_dict[
                            field["id"]
                        ]

        return data_dict


@method_decorator(staff_or_finance_or_contracting_required, name="dispatch")
class ProjectDetailDownloadView(SingleObjectMixin, View):
    model = Project

    def get(self, request, *args, **kwargs):
        export_type = kwargs.get("export_type", None)
        self.object = self.get_object()
        context = self.get_paf_download_context()
        template_path = "application_projects/paf_export.html"

        if export_type == "pdf":
            pdf_page_settings = PDFPageSettings.load(request_or_site=request)
            context["pagesize"] = pdf_page_settings.download_page_size
            context["show_footer"] = True
            return render_as_pdf(
                request=request,
                context=context,
                template_name=template_path,
                filename=self.get_slugified_file_name(export_type),
            )
        elif export_type == "docx":
            context["show_footer"] = False

            return self.render_as_docx(
                context=context,
                template=get_template(template_path),
                filename=self.get_slugified_file_name(export_type),
            )
        else:
            raise Http404(f"{export_type} type not supported at the moment")

    def render_as_docx(self, context, template, filename):
        html = template.render(context)

        buf = io.BytesIO()
        document = Document()
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html, document)
        document.save(buf)

        response = HttpResponse(buf.getvalue(), content_type="application/docx")
        response["Content-Disposition"] = f"attachment; filename={filename}"
        return response

    def get_slugified_file_name(self, export_type):
        return f"{datetime.date.today().strftime('%Y%m%d')}-{slugify(self.object.title)}.{export_type}"

    def get_paf_download_context(self):
        context = {}
        context["id"] = self.object.id
        context["title"] = self.object.title
        context["project_link"] = self.request.build_absolute_uri(
            self.object.get_absolute_url()
        )
        context["contractor_name"] = self.object.user

        context["approvals"] = self.object.paf_approvals.all()
        context["paf_data"] = self.get_paf_data_with_field(self.object)
        context["submission"] = self.object.submission
        context["submission_link"] = self.request.build_absolute_uri(
            reverse(
                "apply:submissions:detail", kwargs={"pk": self.object.submission.id}
            )
        )
        context["supporting_documents"] = self.get_supporting_documents(self.object)
        context["org_name"] = settings.ORG_LONG_NAME
        return context

    def get_paf_data_with_field(self, project):
        data_dict = {}
        form_data_dict = project.form_data
        for field in project.form_fields.raw_data:
            if field.get("type", None) in ["file", "multi_file"]:
                continue
            if field["id"] in form_data_dict.keys():
                if isinstance(field["value"], dict) and "field_label" in field["value"]:
                    data_dict[field["value"]["field_label"]] = form_data_dict[
                        field["id"]
                    ]

        return data_dict

    def get_supporting_documents(self, project):
        documents_dict = {}
        for packet_file in project.packet_files.all():
            documents_dict[packet_file.title] = self.request.build_absolute_uri(
                reverse(
                    "apply:projects:document",
                    kwargs={"pk": project.submission.id, "file_pk": packet_file.id},
                )
            )
        return documents_dict


class ProjectFormsEditView(BaseStreamForm, UpdateView):
    model = Project

    def buttons(self):
        yield ("submit", "primary", _("Save"))

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()

        permission, msg = has_permission(
            "paf_edit", self.request.user, self.object, raise_exception=True
        )
        if not permission:
            messages.info(self.request, msg)
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        submission_attachments = []
        for _field, files in self.object.submission.extract_files().items():
            if isinstance(files, list):
                submission_attachments.extend(files)
            else:
                submission_attachments.append(files)

        return {
            "title": self.object.title,
            "buttons": self.buttons(),
            "object": self.object,
            "submissions_attachments": submission_attachments,
            **kwargs,
        }

    def get_form_class(self, form_class, draft=False, form_data=None, user=None):
        return type(
            "WagtailStreamForm",
            (form_class,),
            self.get_form_fields(draft, form_data, user),
        )


class ProjectFormEditView(ProjectFormsEditView):
    pf_form = None

    template_name = "application_projects/project_approval_form.html"

    @cached_property
    def approval_form(self):
        """Fetch the project form from the fund directly instead of going through round"""
        approval_form = self.object.submission.page.specific.approval_forms.first()

        return approval_form

    def get_pf_form(self, form_class=None):
        if form_class is None:
            form_class = self.get_form_class(ProjectForm)

        return form_class(**self.get_pf_form_kwargs())

    def get_pf_form_kwargs(self):
        kwargs = super().get_form_kwargs()

        if self.approval_form:
            fields = self.approval_form.form.get_form_fields()
        else:
            fields = {}

        kwargs["extra_fields"] = fields
        initial = self.object.raw_data
        for field_id in self.object.file_field_ids:
            initial.pop(field_id + "-uploads", False)
            initial[field_id] = get_placeholder_file(self.object.raw_data.get(field_id))
        kwargs["initial"].update(initial)
        return kwargs

    def get_pf_form_fields(self):
        return self.object.form_fields or self.approval_form.form.form_fields

    def get_context_data(self, **kwargs):
        self.pf_form = self.get_pf_form()

        submission_attachments = []
        for _field, files in self.object.submission.extract_files().items():
            if isinstance(files, list):
                submission_attachments.extend(files)
            else:
                submission_attachments.append(files)

        ctx = {
            "approval_form_exists": True if self.approval_form else False,
            "pf_form": self.pf_form,
            **super().get_context_data(),
            **kwargs,
        }

        return ctx

    def get_defined_fields(self):
        approval_form = self.approval_form
        if approval_form and not self.pf_form:
            return self.get_pf_form_fields()
        return self.object.get_defined_fields()

    def post(self, request, *args, **kwargs):
        """
        Handle POST requests: instantiate a form instance with the passed
        POST variables and then check if it's valid.
        """

        self.pf_form = self.get_pf_form()
        if self.pf_form.is_valid():
            try:
                pf_form_fields = self.get_pf_form_fields()
            except AttributeError:
                pf_form_fields = []
            self.pf_form.save(pf_form_fields=pf_form_fields)
            self.pf_form.delete_temporary_files()
            # remove PAF addition task for staff group
            remove_tasks_for_user(
                code=PROJECT_WAITING_PF,
                user=self.object.lead,
                related_obj=self.object,
            )
            # add project forms submission task for staff group if SOW has been updated
            # OR if the SOW doesn't exist (user_has_updated_sow_details == None)
            if (
                updated_sow := self.object.user_has_updated_sow_details
            ) or updated_sow is None:
                add_task_to_user(
                    code=PROJECT_SUBMIT_PAF,
                    user=self.object.lead,
                    related_obj=self.object,
                )
            return HttpResponseRedirect(self.get_success_url())
        else:
            return self.form_invalid(self.pf_form)


class ProjectSOWEditView(ProjectFormEditView):
    sow_form = None

    template_name = "application_projects/project_sow_form.html"

    @cached_property
    def approval_sow_form(self):
        """Fetch the project form from the fund directly instead of going through round"""
        approval_sow_form = self.object.submission.page.specific.sow_forms.first()

        return approval_sow_form

    def get_sow_form(self, form_class=None):
        if form_class is None:
            form_class = self.get_form_class(ProjectSOWForm)

        return form_class(**self.get_sow_form_kwargs())

    def get_sow_form_fields(self):
        if hasattr(self.object, "sow"):
            return (
                self.object.sow.form_fields or self.approval_sow_form.form.form_fields
            )
        return self.approval_sow_form.form.form_fields

    def get_context_data(self, **kwargs):
        if self.approval_sow_form:
            self.sow_form = self.get_sow_form()

        submission_attachments = []
        for _field, files in self.object.submission.extract_files().items():
            if isinstance(files, list):
                submission_attachments.extend(files)
            else:
                submission_attachments.append(files)

        return {
            "sow_form_exists": True if self.approval_sow_form else False,
            "sow_form": self.sow_form,
            **super().get_context_data(),
            **kwargs,
        }

    def get_sow_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.approval_sow_form:
            fields = self.approval_sow_form.form.get_form_fields()

            kwargs["extra_fields"] = fields
            try:
                sow_instance = self.object.sow
                initial = sow_instance.raw_data
                for field_id in sow_instance.file_field_ids:
                    initial.pop(field_id + "-uploads", False)
                    initial[field_id] = get_placeholder_file(
                        sow_instance.raw_data.get(field_id)
                    )
                initial["project"] = self.object
                kwargs["initial"].update(initial)
            except ObjectDoesNotExist:
                kwargs["initial"].update({"project": self.object})
        return kwargs

    def get_defined_fields(self):
        if self.approval_sow_form and not self.sow_form:
            return self.get_sow_form_fields()
        return self.object.get_defined_fields()

    def post(self, request, *args, **kwargs):
        """
        Handle POST requests: instantiate a form instance with the passed
        POST variables and then check if it's valid.
        """

        if self.approval_sow_form:
            self.sow_form = self.get_sow_form()
            if self.sow_form.is_valid():
                try:
                    sow_form_fields = self.get_sow_form_fields()
                except AttributeError:
                    sow_form_fields = []

                self.sow_form.save(sow_form_fields=sow_form_fields, project=self.object)
                self.sow_form.delete_temporary_files()
                # remove SOW addition task for staff group
                remove_tasks_for_user(
                    code=PROJECT_WAITING_SOW,
                    user=self.object.lead,
                    related_obj=self.object,
                )
                # add project forms submission task for staff group if project form has been updated
                if self.object.user_has_updated_pf_details:
                    add_task_to_user(
                        code=PROJECT_SUBMIT_PAF,
                        user=self.object.lead,
                        related_obj=self.object,
                    )
                return HttpResponseRedirect(self.get_success_url())
            else:
                return self.form_invalid(self.sow_form)


@method_decorator(staff_or_finance_or_contracting_required, name="dispatch")
class ProjectListView(SingleTableMixin, FilterView):
    filterset_class = ProjectListFilter
    queryset = Project.objects.for_table()
    table_class = ProjectsListTable
    template_name = "application_projects/project_list.html"
